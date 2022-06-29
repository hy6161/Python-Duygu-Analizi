import docx
from textblob import TextBlob
import pandas as pd
from docx import Document
from docx.shared import Inches
document = Document()
doc = docx.Document("demo.docx")
all_paras = doc.paragraphs

metin=[]
for para in all_paras:
    metin.append(para.text)

print(metin)
metinislem=str(metin[0]).split(".")
metintest=[]
for i in range(len(metinislem)-1):
    metintest.append(metinislem[i])

print(metintest)

olumlu_yazılar=[]
olumlu_sonuc=[]
olumsuz_yazılar=[]
olumsuz_sonuc=[]
yorumsuz=[]

for yazı in metintest:
    blob1=TextBlob(yazı)
    blob_eng=blob1.translate(from_lang="tr",to = "en")

    if(blob_eng.polarity>0):
        olumlu_yazılar.append(yazı)
        olumlu_sonuc.append(blob_eng.sentiment)
    elif(blob_eng.polarity<0):
        olumsuz_yazılar.append(yazı)
        olumsuz_sonuc.append(blob_eng.sentiment)
    else:
        yorumsuz.append(yazı)



document.add_heading('Duygu Analizi', 0)
document.add_heading('Olumlu Yazılar', level=1)
document.add_paragraph(f"{olumlu_yazılar}\n")
document.add_heading('Olumsuz Yazılar', level=1)
document.add_paragraph(olumsuz_yazılar)
document.add_heading('Yorumsuz Yazılar', level=1)
document.add_paragraph(yorumsuz)

document.save('sonuç.docx')